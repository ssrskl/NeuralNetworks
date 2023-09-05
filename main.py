from docx import Document
# 创建一个文档对象
document = Document('demo.docx')
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
document.add_heading("Hello World",level=1)
document.sections[0].header.paragraphs[0].text = "这是第1节页眉"
document.save('demo.docx')
